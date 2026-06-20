from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports" / "dsr-webstack-progress-report.docx"
PDF_OUT = ROOT / "reports" / "dsr-webstack-progress-report.pdf"


TOKENS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "muted": "555555",
    "border": "DADCE0",
    "header_fill": "F2F4F7",
    "success_fill": "EAF6EE",
    "note_fill": "F4F6F9",
}


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def shade_header(row):
    for cell in row.cells:
        set_cell_fill(cell, TOKENS["header_fill"])
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(11, 37, 69)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    shade_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = str(value)
    set_table_geometry(table, widths)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.10
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.5)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, TOKENS["note_fill"])
    set_table_geometry(table, [9360])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p.add_run(" " + body)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, TOKENS["blue"], 16, 8),
        ("Heading 2", 13, TOKENS["blue"], 12, 6),
        ("Heading 3", 12, TOKENS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("DSR Portal Progress Report").font.size = Pt(9)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("DSR Portal Webstack Progress Report")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("Architecture, implementation status, cleanup summary, and next milestones")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", "updated-webstack-of-dsr--portal"],
            ["Prepared on", "11 June 2026"],
            ["Scope", "Full webstack: Next.js web app, Express API, Prisma/PostgreSQL schema, and legacy DSR portal frontend"],
            ["Report status", "Working progress report based on current local repository state"],
        ],
        [2200, 7160],
    )
    doc.add_page_break()


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=22,
            leading=26,
            textColor=colors.HexColor("#0B2545"),
            alignment=0,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportSubtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            textColor=colors.HexColor("#555555"),
            spaceAfter=14,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Report",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodyReport",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=9.6,
            leading=13,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SmallReport",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11,
        )
    )
    return styles


def pdf_table(data, widths, styles):
    body = []
    for row in data:
        body.append([Paragraph(str(cell), styles["SmallReport"]) for cell in row])
    table = Table(body, colWidths=widths, hAlign="LEFT", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DADCE0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf_report():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title="DSR Portal Webstack Progress Report",
    )
    story = [
        Paragraph("DSR Portal Webstack Progress Report", styles["ReportTitle"]),
        Paragraph(
            "Architecture, implementation status, cleanup summary, and next milestones",
            styles["ReportSubtitle"],
        ),
        pdf_table(
            [
                ["Field", "Value"],
                ["Project", "updated-webstack-of-dsr--portal"],
                ["Prepared on", "11 June 2026"],
                ["Scope", "Full webstack: Next.js web app, Express API, Prisma/PostgreSQL schema, and legacy DSR portal frontend"],
                ["Report status", "Working progress report based on current local repository state"],
            ],
            [1.35 * inch, 5.0 * inch],
            styles,
        ),
        PageBreak(),
        Paragraph("Executive Summary", styles["H1Report"]),
        Paragraph(
            "The DSR Portal repository is a modernized monorepo around a Next.js web package, an Express API package, "
            "a Prisma/PostgreSQL persistence layer, and a large legacy browser portal that remains the main feature surface "
            "for DSR data entry, annexures, PDF workflows, dashboard tracking, roles, and project phases.",
            styles["BodyReport"],
        ),
        Paragraph(
            "<b>Current direction:</b> Keep the working project stable while cleaning visible background effects, redundant temporary files, "
            "and low-risk noise. Large CSS override cleanup should be handled in reviewed batches because many duplicate-looking rules are cross-file fallbacks.",
            styles["BodyReport"],
        ),
        Paragraph("Webstack Overview", styles["H1Report"]),
        pdf_table(
            [
                ["Layer", "Current implementation", "Purpose"],
                ["Workspace", "npm workspaces: apps/* and packages/*", "Keeps API and web packages under one project root."],
                ["Web", "Next.js 15, React 19, TypeScript, Tailwind-related utilities", "Modern app shell and local dev server on port 3000."],
                ["Legacy UI", "Static HTML/CSS/JS bundle under apps/web/public/legacy", "Current DSR portal screens, dashboard, annexures, PDF preview, auth UI, and theme system."],
                ["API", "Express 4, TypeScript, helmet, cors, rate limits, cookie-parser, JWT", "Authenticated backend routes for projects, files, reports, users, dashboard, and PDF actions."],
                ["Database", "Prisma 6 with PostgreSQL datasource", "Users, projects, reports, workflow history, audit logs, and uploaded DSR files."],
                ["Jobs/files", "BullMQ, multer, AWS S3 SDK", "Background job and file upload/storage foundation."],
            ],
            [1.15 * inch, 2.2 * inch, 3.0 * inch],
            styles,
        ),
        Paragraph("Repository Metrics", styles["H1Report"]),
        pdf_table(
            [
                ["Metric", "Value", "Notes"],
                ["Tracked source files counted", "153", "Excludes node_modules, .git, package-lock.json, and large legacy projects.json demo data."],
                ["Approximate source lines", "39k-43k", "Range depends on whether generated HTML and binary-like assets are counted."],
                ["Main legacy JS files", "42", "Annexures, project state, PDF preview, navigation, roles, auth, theme, and dashboard behavior."],
                ["Main legacy HTML templates/files", "46", "Build script composes templates into login.html, home.html, and index.html."],
                ["Main CSS files", "3", "styles.css, premium-theme.css, and supporting style assets."],
            ],
            [2.0 * inch, 1.2 * inch, 3.15 * inch],
            styles,
        ),
        Paragraph("Implemented Functional Areas", styles["H1Report"]),
    ]
    for item in [
        "Role-aware login and portal access paths for staff, government authority, and SDLC style entry points.",
        "Project dashboard with district filtering, project cards, progress/status indicators, and shortcut search.",
        "Project lifecycle storage including active project state, phase metadata, phase locks, and child phase creation.",
        "Front matter, chapters, plates, annexures, more-annexures, tables, graphs, signatures, workflow, and audit views.",
        "PDF upload, preview, generated final PDF download, and annexure-specific PDF generation flows.",
        "Backend authentication, dashboard stats, project CRUD/state persistence, report workflow, file upload, jobs, users, and PDF endpoints.",
        "Dark/light theme handling with visual branding, watermark, and legacy app styling continuity.",
    ]:
        story.append(Paragraph("- " + item, styles["BodyReport"]))
    story.extend(
        [
            Paragraph("Recent Progress Completed", styles["H1Report"]),
            pdf_table(
                [
                    ["Area", "Progress", "Project impact"],
                    ["Background system", "Removed old DOM-based live-lines overlay/classes and replaced with CSS-only global grid plus four floating directional lines.", "Lower DOM noise; same visible intent; no feature logic changed."],
                    ["Dark mode visible lines", "Removed the earlier harsh neon/drop-shadow line traces from the old system.", "Dark mode background is calmer and less intrusive."],
                    ["Templates/generated HTML", "Regenerated login.html, home.html, and index.html from templates with asset version grid-live-lines-cleanup-20260611.", "Browser cache receives the corrected CSS/JS version."],
                    ["Cleanup", "Deleted fix.js, an unused one-off encoding repair script with no app references.", "Safe removal; not part of runtime."],
                    ["Console noise", "Removed harmless informational console logs from frontmatter/main paths while keeping user-facing fallbacks.", "Cleaner dev console without changing UI behavior."],
                    ["CSS safety", "Restored malformed reset selectors and converted an orphan heading into a comment.", "CSS parses cleanly and avoids accidental invalid selector behavior."],
                ],
                [1.35 * inch, 3.0 * inch, 2.0 * inch],
                styles,
            ),
            Paragraph("Current Verification Status", styles["H1Report"]),
            pdf_table(
                [
                    ["Check", "Result"],
                    ["CSS parse check", "styles.css OK; premium-theme.css OK."],
                    ["Old line system search", "No live-lines-overlay, interactive-grid-bg, live-line, old scan animation, dense/neon line variant, or old drop-shadow trace found."],
                    ["Scoped diff for cleanup", "11 legacy files changed in the cleanup scope: 179 insertions and 2,338 deletions, mostly comments/blank lines plus old line system removal."],
                    ["Risk posture", "No aggressive CSS override deletion performed beyond low-risk exact/unused items."],
                ],
                [2.25 * inch, 4.1 * inch],
                styles,
            ),
            Paragraph("Backend/Data Model Progress", styles["H1Report"]),
            pdf_table(
                [
                    ["Component", "Status"],
                    ["Prisma enums", "Role, ProjectStatus, and ReportStatus are present."],
                    ["Core models", "User, Project, Report, WorkflowHistory, AuditLog, and DsrFile support the portal workflow."],
                    ["Phase model additions", "Project includes phaseNo, parentPhaseId, phaseLocked, phaseOrigin, and parent/child phase relations."],
                    ["Main API routes", "Auth, dashboard, files, jobs, projects, reports, users, and PDF endpoints are wired in server.ts."],
                    ["Security middleware", "Helmet, CORS, cookie parser, JSON body limits, route-level auth, audit mutations, and rate limits are present."],
                ],
                [2.05 * inch, 4.3 * inch],
                styles,
            ),
            Paragraph("Cleanup Candidates and Risk", styles["H1Report"]),
            Paragraph(
                "A CSS audit found many candidates that look overwritten or duplicated, but not all are safe to remove immediately. "
                "The legacy app loads styles differently across login, generated portal pages, and home page, so a rule that appears redundant "
                "in one page can still act as fallback in another.",
                styles["BodyReport"],
            ),
            pdf_table(
                [
                    ["Candidate", "Observed count", "Recommendation"],
                    ["Exact duplicate CSS candidates", "About 52 after repair audit", "Remove only when selector/file/page context is identical."],
                    ["Overwritten CSS candidates", "About 561", "Batch by component: auth, topbar, sidebar, dashboard, annexures, PDF preview."],
                    ["Important declarations", "About 1,876", "Do not bulk-remove; many are legacy override contracts."],
                    ["Large demo/project data", "projects.json excluded from LOC count", "Review separately before deleting because it may preserve demo state."],
                ],
                [2.1 * inch, 1.35 * inch, 2.9 * inch],
                styles,
            ),
            Paragraph("Recommended Next Milestones", styles["H1Report"]),
        ]
    )
    for idx, item in enumerate(
        [
            "Create a visual regression checklist for login, home, dashboard, projects, workflow, front matter, annexures, PDF preview, and dark mode.",
            "Split CSS cleanup into small batches and verify each batch in browser before continuing.",
            "Move generated HTML handling behind the existing build script discipline so templates remain the source of truth.",
            "Add API smoke tests for auth, projects, project state save/load, phase creation, and PDF endpoints.",
            "Document the role/permission matrix so future UI cleanup does not accidentally expose restricted actions.",
            "Decide whether legacy UI remains the primary UI or is gradually migrated into Next.js routes.",
        ],
        start=1,
    ):
        story.append(Paragraph(f"{idx}. {item}", styles["BodyReport"]))
    story.extend(
        [
            Paragraph("Conclusion", styles["H1Report"]),
            Paragraph(
                "The webstack is functional and broad: it already covers the key DSR portal workflows from authentication and project creation "
                "through annexure editing, PDF handling, workflow tracking, and backend persistence. The safest improvement path is incremental: "
                "keep runtime behavior stable, remove unused temporary artifacts, and treat CSS override cleanup as a measured refactor with visual checks.",
                styles["BodyReport"],
            ),
        ]
    )
    doc.build(story)
    print(PDF_OUT)


def build_report():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The DSR Portal repository is a modernized monorepo around a Next.js web package, an Express API package, "
        "a Prisma/PostgreSQL persistence layer, and a large legacy browser portal that remains the main feature surface "
        "for DSR data entry, annexures, PDF workflows, dashboard tracking, roles, and project phases."
    )
    add_note(
        doc,
        "Current direction:",
        "Keep the working project stable while cleaning visible background effects, redundant temporary files, and low-risk noise. "
        "Large CSS override cleanup should be handled in reviewed batches because many duplicate-looking rules are cross-file fallbacks.",
    )

    doc.add_heading("Webstack Overview", level=1)
    add_table(
        doc,
        ["Layer", "Current implementation", "Purpose"],
        [
            ["Workspace", "npm workspaces: apps/* and packages/*", "Keeps API and web packages under one project root."],
            ["Web", "Next.js 15, React 19, TypeScript, Tailwind-related utilities", "Modern app shell and local dev server on port 3000."],
            ["Legacy UI", "Static HTML/CSS/JS bundle under apps/web/public/legacy", "Current DSR portal screens, dashboard, annexures, PDF preview, auth UI, and theme system."],
            ["API", "Express 4, TypeScript, helmet, cors, rate limits, cookie-parser, JWT", "Authenticated backend routes for projects, files, reports, users, dashboard, and PDF actions."],
            ["Database", "Prisma 6 with PostgreSQL datasource", "Users, projects, reports, workflow history, audit logs, and uploaded DSR files."],
            ["Jobs/files", "BullMQ, multer, AWS S3 SDK", "Background job and file upload/storage foundation."],
        ],
        [1700, 3300, 4360],
    )

    doc.add_heading("Repository Metrics", level=1)
    add_table(
        doc,
        ["Metric", "Value", "Notes"],
        [
            ["Tracked source files counted", "153", "Excludes node_modules, .git, package-lock.json, and large legacy projects.json demo data."],
            ["Approximate source lines", "39k-43k", "Range depends on whether generated HTML and binary-like assets are counted."],
            ["Main legacy JS files", "42", "Annexures, project state, PDF preview, navigation, roles, auth, theme, and dashboard behavior."],
            ["Main legacy HTML templates/files", "46", "Build script composes templates into login.html, home.html, and index.html."],
            ["Main CSS files", "3", "styles.css, premium-theme.css, and supporting style assets."],
        ],
        [2600, 2000, 4760],
    )

    doc.add_heading("Implemented Functional Areas", level=1)
    for item in [
        "Role-aware login and portal access paths for staff, government authority, and SDLC style entry points.",
        "Project dashboard with district filtering, project cards, progress/status indicators, and shortcut search.",
        "Project lifecycle storage including active project state, phase metadata, phase locks, and child phase creation.",
        "Front matter, chapters, plates, annexures, more-annexures, tables, graphs, signatures, workflow, and audit views.",
        "PDF upload, preview, generated final PDF download, and annexure-specific PDF generation flows.",
        "Backend authentication, dashboard stats, project CRUD/state persistence, report workflow, file upload, jobs, users, and PDF endpoints.",
        "Dark/light theme handling with visual branding, watermark, and legacy app styling continuity.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Recent Progress Completed", level=1)
    add_table(
        doc,
        ["Area", "Progress", "Project impact"],
        [
            ["Background system", "Removed old DOM-based live-lines overlay/classes and replaced with CSS-only global grid plus four floating directional lines.", "Lower DOM noise; same visible intent; no feature logic changed."],
            ["Dark mode visible lines", "Removed the earlier harsh neon/drop-shadow line traces from the old system.", "Dark mode background is calmer and less intrusive."],
            ["Templates/generated HTML", "Regenerated login.html, home.html, and index.html from templates with asset version grid-live-lines-cleanup-20260611.", "Browser cache receives the corrected CSS/JS version."],
            ["Cleanup", "Deleted fix.js, an unused one-off encoding repair script with no app references.", "Safe removal; not part of runtime."],
            ["Console noise", "Removed harmless informational console logs from frontmatter/main paths while keeping user-facing fallbacks.", "Cleaner dev console without changing UI behavior."],
            ["CSS safety", "Restored malformed reset selectors and converted an orphan heading into a comment.", "CSS parses cleanly and avoids accidental invalid selector behavior."],
        ],
        [2100, 4300, 2960],
    )

    doc.add_heading("Current Verification Status", level=1)
    add_table(
        doc,
        ["Check", "Result"],
        [
            ["CSS parse check", "styles.css OK; premium-theme.css OK."],
            ["Old line system search", "No live-lines-overlay, interactive-grid-bg, live-line, old scan animation, dense/neon line variant, or old drop-shadow trace found."],
            ["Scoped diff for cleanup", "11 legacy files changed in the cleanup scope: 179 insertions and 2,338 deletions, mostly comments/blank lines plus old line system removal."],
            ["Risk posture", "No aggressive CSS override deletion performed beyond low-risk exact/unused items."],
        ],
        [3300, 6060],
    )

    doc.add_heading("Backend/Data Model Progress", level=1)
    add_table(
        doc,
        ["Component", "Status"],
        [
            ["Prisma enums", "Role, ProjectStatus, and ReportStatus are present."],
            ["Core models", "User, Project, Report, WorkflowHistory, AuditLog, and DsrFile support the portal workflow."],
            ["Phase model additions", "Project includes phaseNo, parentPhaseId, phaseLocked, phaseOrigin, and parent/child phase relations."],
            ["Main API routes", "Auth, dashboard, files, jobs, projects, reports, users, and PDF endpoints are wired in server.ts."],
            ["Security middleware", "Helmet, CORS, cookie parser, JSON body limits, route-level auth, audit mutations, and rate limits are present."],
        ],
        [3000, 6360],
    )

    doc.add_heading("Cleanup Candidates and Risk", level=1)
    doc.add_paragraph(
        "A CSS audit found many candidates that look overwritten or duplicated, but not all are safe to remove immediately. "
        "The legacy app loads styles differently across login, generated portal pages, and home page, so a rule that appears redundant "
        "in one page can still act as fallback in another."
    )
    add_table(
        doc,
        ["Candidate", "Observed count", "Recommendation"],
        [
            ["Exact duplicate CSS candidates", "About 52 after repair audit", "Remove only when selector/file/page context is identical."],
            ["Overwritten CSS candidates", "About 561", "Batch by component: auth, topbar, sidebar, dashboard, annexures, PDF preview."],
            ["Important declarations", "About 1,876", "Do not bulk-remove; many are legacy override contracts."],
            ["Large demo/project data", "projects.json excluded from LOC count", "Review separately before deleting because it may preserve demo state."],
        ],
        [3100, 1800, 4460],
    )

    doc.add_heading("Recommended Next Milestones", level=1)
    for item in [
        "Create a visual regression checklist for login, home, dashboard, projects, workflow, front matter, annexures, PDF preview, and dark mode.",
        "Split CSS cleanup into small batches and verify each batch in browser before continuing.",
        "Move generated HTML handling behind the existing build script discipline so templates remain the source of truth.",
        "Add API smoke tests for auth, projects, project state save/load, phase creation, and PDF endpoints.",
        "Document the role/permission matrix so future UI cleanup does not accidentally expose restricted actions.",
        "Decide whether legacy UI remains the primary UI or is gradually migrated into Next.js routes.",
    ]:
        add_numbered(doc, item)

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "The webstack is functional and broad: it already covers the key DSR portal workflows from authentication and project creation "
        "through annexure editing, PDF handling, workflow tracking, and backend persistence. The safest improvement path is incremental: "
        "keep runtime behavior stable, remove unused temporary artifacts, and treat CSS override cleanup as a measured refactor with visual checks."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    build_pdf_report()
    print(OUT)


if __name__ == "__main__":
    build_report()
